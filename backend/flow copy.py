{
 "cells": [
  {
   "cell_type": "markdown",
   "id": "1eefaa99",
   "metadata": {},
   "source": [
    "Word Document Formatter"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "8c70c933",
   "metadata": {},
   "source": [
    "1. importing libraries"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 35,
   "id": "1c33bffd",
   "metadata": {},
   "outputs": [],
   "source": [
    "from langgraph.graph import StateGraph\n",
    "from langchain_core.messages import BaseMessage, HumanMessage\n",
    "from langchain_openai import ChatOpenAI\n",
    "from langchain_core.prompts import PromptTemplate\n",
    "from typing import TypedDict, Annotated, List, Dict, Optional, Literal\n",
    "from langgraph.graph import add_messages\n",
    "from langchain_core.output_parsers import StrOutputParser\n",
    "from pydantic import BaseModel, Field\n",
    "from langchain_core.output_parsers import PydanticOutputParser"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 36,
   "id": "9a851d41",
   "metadata": {},
   "outputs": [
    {
     "data": {
      "text/plain": [
       "False"
      ]
     },
     "execution_count": 36,
     "metadata": {},
     "output_type": "execute_result"
    }
   ],
   "source": [
    "from dotenv import load_dotenv\n",
    "load_dotenv()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 37,
   "id": "da796606",
   "metadata": {},
   "outputs": [],
   "source": [
    "model = ChatOpenAI(\n",
    "    model=\"mistralai/devstral-2512:free\",\n",
    "    openai_api_base=\"https://openrouter.ai/api/v1\",\n",
    "    openai_api_key=\"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\",\n",
    ")"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "742e713d",
   "metadata": {},
   "source": [
    "2.states define"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 38,
   "id": "9544c3d5",
   "metadata": {},
   "outputs": [],
   "source": [
    "class DocStructure(BaseModel):\n",
    "    document_type: str = Field(\n",
    "        description=\"Type of document: academic, research, business, general\"\n",
    "    )\n",
    "\n",
    "    headings: Dict[str, int] = Field(\n",
    "        description=\"Detected headings mapped to heading level (e.g., {'Introduction': 1})\"\n",
    "    )\n",
    "\n",
    "    body_font_size: int = Field(\n",
    "        description=\"Target font size for body text\"\n",
    "    )\n",
    "\n",
    "    heading_font_size: int = Field(\n",
    "        description=\"Target font size for headings\"\n",
    "    )\n",
    "\n",
    "    alignment: str = Field(\n",
    "        description=\"Text alignment: left, right, center, justify\"\n",
    "    )\n",
    "\n",
    "    issues_found: List[str] = Field(\n",
    "        description=\"Formatting or structural issues detected\"\n",
    "    )\n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 39,
   "id": "bc7cc7ac",
   "metadata": {},
   "outputs": [],
   "source": [
    "class FormatCommand(BaseModel):\n",
    "    action: str\n",
    "    target: str\n",
    "    value: int | str"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 40,
   "id": "33f1d910",
   "metadata": {},
   "outputs": [],
   "source": [
    "class wordFormatter(TypedDict, total=False):\n",
    "    doc: str\n",
    "    structDoc: Annotated[list, lambda x, y: y]\n",
    "    formatPlan: List[DocStructure]\n",
    "    formatCmd: List[FormatCommand]\n",
    "    output_doc: Optional[str]\n",
    "    output_pdf: Optional[str]\n",
    "    user_instruction: str"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "732e9985",
   "metadata": {},
   "source": [
    "3. nodes Building"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "c9f4a615",
   "metadata": {},
   "source": [
    "NodeFn1: Loading the document"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 41,
   "id": "2e791913",
   "metadata": {},
   "outputs": [],
   "source": [
    "from docx import Document\n",
    "\n",
    "def load_docx(path: str):\n",
    "    doc = Document(path)\n",
    "    paragraphs = []\n",
    "\n",
    "    for p in doc.paragraphs:\n",
    "        paragraphs.append({\n",
    "            \"text\": p.text.strip(),\n",
    "            \"style\": p.style.name if p.style else None,\n",
    "            \"alignment\": p.alignment,\n",
    "            \"runs\": [\n",
    "                {\n",
    "                    \"text\": r.text,\n",
    "                    \"bold\": r.bold,\n",
    "                    \"italic\": r.italic,\n",
    "                    \"font_size\": r.font.size.pt if r.font.size else None\n",
    "                }\n",
    "                for r in p.runs\n",
    "            ]\n",
    "        })\n",
    "    return paragraphs\n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 42,
   "id": "4bb38dae",
   "metadata": {},
   "outputs": [
    {
     "name": "stdout",
     "output_type": "stream",
     "text": [
      "[{'text': 'Introduction', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Introduction', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Qwertyujnbvcddx', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'Qwertyujnbvcddx', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Ertyhhbvcx', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'Ertyhhbvcx', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'sdfgbhcd', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'sdfgbhcd', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'This document is created for a word processing formatting project. The purpose of this file is to allow the user to apply formatting such as text justification, font size, font color, line spacing, and paragraph alignment.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'This ', 'bold': None, 'italic': None, 'font_size': None}, {'text': 'document ', 'bold': None, 'italic': None, 'font_size': 18.0}, {'text': 'is', 'bold': None, 'italic': None, 'font_size': 18.0}, {'text': ' created for a word processing formatting project. The purpose of this file is to allow the user to apply formatting such as text justification, font size, font color, line spacing, and paragraph alignment.', 'bold': None, 'italic': None, 'font_size': 18.0}]}, {'text': 'Project Description', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Project Description', 'bold': None, 'italic': None, 'font_size': 10.0}]}, {'text': 'The project focuses on understanding basic document formatting techniques commonly used in professional and academic documents. Students are expected to modify the appearance of the text without changing the actual content provided in this document.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'The project focuses on understanding basic document formatting techniques commonly used in professional and academic documents. Students are expected to modify the appearance of the text without changing the actual content provided in this document.', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Key points:', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'Key points:', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To handle abc', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To handle ', 'bold': None, 'italic': None, 'font_size': None}, {'text': 'abc', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To handle research', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To handle research', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To ease assignment', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To ease assignment', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Conclusion', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Conclusion', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'By applying appropriate formatting styles, this document can be transformed into a well-structured and visually appealing report. Proper formatting improves readability and enhances the overall presentation of written content.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'By applying appropriate formatting styles, this document can be transformed into a well-structured and visually appealing report. Proper formatting improves readability and enhances the overall presentation of written content.', 'bold': None, 'italic': None, 'font_size': 14.0}]}]\n"
     ]
    }
   ],
   "source": [
    "loader = load_docx('Draft.docx')\n",
    "# doc = loader.load()\n",
    "print(loader)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 43,
   "id": "557b6928",
   "metadata": {},
   "outputs": [],
   "source": [
    "def loadDocNode(state: wordFormatter) -> wordFormatter:\n",
    "    doc = state['doc']\n",
    "    state['structDoc'] = load_docx(doc)\n",
    "    return state"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "d61ee019",
   "metadata": {},
   "source": [
    "NodeFn2: Structure Analysis Agent node"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 44,
   "id": "63d0bbcf",
   "metadata": {},
   "outputs": [],
   "source": [
    "def docStructureNode(state: wordFormatter) -> wordFormatter:\n",
    "    struct_doc = state[\"structDoc\"]\n",
    "\n",
    "    parser = PydanticOutputParser(pydantic_object=DocStructure)\n",
    "\n",
    "    prompt = PromptTemplate(\n",
    "        template=\"\"\"\n",
    "You are a document structure analysis agent.\n",
    "\n",
    "Analyze the following document structure and produce a formatting plan.\n",
    "\n",
    "Document structure:\n",
    "{struct_doc}\n",
    "\n",
    "{format_instructions}\n",
    "\"\"\",\n",
    "        input_variables=[\"struct_doc\"],\n",
    "        partial_variables={\n",
    "            \"format_instructions\": parser.get_format_instructions()\n",
    "        }\n",
    "    )\n",
    "\n",
    "    response = model.invoke(\n",
    "        prompt.format(struct_doc=struct_doc)\n",
    "    )\n",
    "\n",
    "    state[\"formatPlan\"] = parser.parse(response.content)\n",
    "    return state"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "3e9479af",
   "metadata": {},
   "source": [
    "NodeFn3: Formatter Agent node"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "0caf77ea",
   "metadata": {},
   "source": [
    "3.1 word formatter"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 45,
   "id": "91fc7e77",
   "metadata": {},
   "outputs": [],
   "source": [
    "from docx import Document\n",
    "from docx.enum.text import WD_ALIGN_PARAGRAPH\n",
    "from docx.shared import Pt\n",
    "\n",
    "def formatterNode(state: wordFormatter) -> wordFormatter:\n",
    "    plan = state[\"formatPlan\"]\n",
    "    doc_path = state[\"doc\"]\n",
    "\n",
    "    doc = Document(doc_path)\n",
    "\n",
    "    for p in doc.paragraphs:\n",
    "        # Headings\n",
    "        if p.style.name.startswith(\"Heading\"):\n",
    "            for run in p.runs:\n",
    "                run.font.size = Pt(plan.heading_font_size)\n",
    "        else:\n",
    "            # Body text\n",
    "            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY\n",
    "            for run in p.runs:\n",
    "                run.font.size = Pt(plan.body_font_size)\n",
    "\n",
    "    output_path = \"formatted.docx\"\n",
    "    doc.save(output_path)\n",
    "\n",
    "    state[\"output_doc\"] = output_path\n",
    "    return state"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "id": "4925d8d5",
   "metadata": {},
   "outputs": [],
   "source": []
  },
  {
   "cell_type": "markdown",
   "id": "dc9cbcc3",
   "metadata": {},
   "source": [
    "3.2 pdf generater"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 46,
   "id": "6b12b384",
   "metadata": {},
   "outputs": [],
   "source": [
    "import subprocess\n",
    "\n",
    "def pdfGeneratorNode(state: wordFormatter):\n",
    "    docx_path = state[\"output_doc\"]\n",
    "    pdf_path = docx_path.replace(\".docx\", \".pdf\")\n",
    "\n",
    "    subprocess.run([\n",
    "        \"libreoffice\",\n",
    "        \"--headless\",\n",
    "        \"--convert-to\", \"pdf\",\n",
    "        docx_path,\n",
    "        \"--outdir\", \".\"\n",
    "    ], check=True)\n",
    "\n",
    "    return {\"output_pdf\": pdf_path}\n"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "56ea64a9",
   "metadata": {},
   "source": [
    "NodeFn4: User Instructions node"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "id": "2c79ada6",
   "metadata": {},
   "outputs": [],
   "source": []
  },
  {
   "cell_type": "code",
   "execution_count": 47,
   "id": "71310190",
   "metadata": {},
   "outputs": [],
   "source": [
    "def chatCommandNode(state: wordFormatter):\n",
    "    instruction = state[\"user_instruction\"]\n",
    "\n",
    "    prompt = PromptTemplate(\n",
    "        template=\"\"\"\n",
    "        Convert the user instruction into a formatting command.\n",
    "        Instruction: {instruction}\n",
    "        \"\"\",\n",
    "        input_variables=[\"instruction\"]\n",
    "    )\n",
    "\n",
    "    command = model.with_structured_output(FormatCommand).invoke(prompt)\n",
    "    state[\"formatCmd\"] = command\n",
    "    return state"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "42a3ecc4",
   "metadata": {},
   "source": [
    "NodeFn5: Chat Command Formatter Node"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 48,
   "id": "d4fc5dd3",
   "metadata": {},
   "outputs": [],
   "source": [
    "from docx.enum.text import WD_ALIGN_PARAGRAPH\n",
    "\n",
    "from docx import Document\n",
    "from docx.shared import Pt\n",
    "from docx.enum.text import WD_ALIGN_PARAGRAPH\n",
    "\n",
    "def applyChatFormatNode(state: wordFormatter) -> wordFormatter:\n",
    "    doc = Document(state[\"doc\"])\n",
    "    instruction = state.get(\"user_instruction\", \"\").lower()\n",
    "\n",
    "    for p in doc.paragraphs:\n",
    "        if \"justify\" in instruction:\n",
    "            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY\n",
    "\n",
    "        if \"font size 12\" in instruction:\n",
    "            for run in p.runs:\n",
    "                run.font.size = Pt(12)\n",
    "\n",
    "    output_path = \"formatted.docx\"\n",
    "    doc.save(output_path)\n",
    "\n",
    "    state[\"output_doc\"] = output_path\n",
    "    return state\n",
    "\n"
   ]
  },
  {
   "cell_type": "markdown",
   "id": "09025d4c",
   "metadata": {},
   "source": [
    "Conditinal path define"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 49,
   "id": "741236ce",
   "metadata": {},
   "outputs": [],
   "source": [
    "def route_after_load(state: wordFormatter):\n",
    "    if state.get(\"user_instruction\"):\n",
    "        return \"chat_command\"\n",
    "    return \"doc_structure\""
   ]
  },
  {
   "cell_type": "markdown",
   "id": "f4d6cba8",
   "metadata": {},
   "source": [
    "4.Graph Struct"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 50,
   "id": "19485487",
   "metadata": {},
   "outputs": [],
   "source": [
    "graph = StateGraph(wordFormatter)\n",
    "\n",
    "graph.add_node(\"load_docx\", loadDocNode)\n",
    "graph.add_node(\"doc_structure\", docStructureNode)\n",
    "graph.add_node(\"doc_formatter\", formatterNode)\n",
    "graph.add_node(\"doc_pdf\", pdfGeneratorNode)\n",
    "graph.add_node(\"chat_command\", chatCommandNode)\n",
    "graph.add_node(\"chat_format\", applyChatFormatNode)\n",
    "\n",
    "graph.set_entry_point(\"load_docx\")\n",
    "graph.add_conditional_edges(\"load_docx\",route_after_load,{\n",
    "        \"chat_command\": \"chat_command\",\n",
    "        \"doc_structure\": \"doc_structure\"\n",
    "    }\n",
    ")\n",
    "\n",
    "graph.add_edge(\"doc_structure\", \"doc_formatter\")\n",
    "\n",
    "graph.add_edge(\"chat_command\", \"chat_format\")\n",
    "\n",
    "# graph.add_edge(\"doc_formatter\", \"doc_pdf\")\n",
    "# graph.add_edge(\"chat_format\", \"doc_pdf\")\n",
    "graph.set_finish_point(\"chat_format\")\n",
    "\n",
    "app = graph.compile()\n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 51,
   "id": "b2392ca8",
   "metadata": {},
   "outputs": [
    {
     "name": "stdout",
     "output_type": "stream",
     "text": [
      "               +-----------+                  \n",
      "               | __start__ |                  \n",
      "               +-----------+                  \n",
      "                      *                       \n",
      "                      *                       \n",
      "                      *                       \n",
      "               +-----------+                  \n",
      "               | load_docx |                  \n",
      "               +-----------+                  \n",
      "              ..            ..                \n",
      "            ..                ..              \n",
      "          ..                    ..            \n",
      "+--------------+           +---------------+  \n",
      "| chat_command |           | doc_structure |  \n",
      "+--------------+           +---------------+  \n",
      "        *                           *         \n",
      "        *                           *         \n",
      "        *                           *         \n",
      "+-------------+            +---------------+  \n",
      "| chat_format |            | doc_formatter |  \n",
      "+-------------+            +---------------+  \n",
      "              **            **                \n",
      "                **        **                  \n",
      "                  **    **                    \n",
      "                +---------+                   \n",
      "                | __end__ |                   \n",
      "                +---------+                   \n"
     ]
    }
   ],
   "source": [
    "app.get_graph().print_ascii()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "id": "ff2ba0e0",
   "metadata": {},
   "outputs": [
    {
     "name": "stdout",
     "output_type": "stream",
     "text": [
      "formatted.docx\n",
      "[{'text': 'Introduction', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Introduction', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Qwertyujnbvcddx', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'Qwertyujnbvcddx', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Ertyhhbvcx', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'Ertyhhbvcx', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'sdfgbhcd', 'style': 'List Paragraph', 'alignment': None, 'runs': [{'text': 'sdfgbhcd', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'This document is created for a word processing formatting project. The purpose of this file is to allow the user to apply formatting such as text justification, font size, font color, line spacing, and paragraph alignment.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'This ', 'bold': None, 'italic': None, 'font_size': None}, {'text': 'document ', 'bold': None, 'italic': None, 'font_size': 18.0}, {'text': 'is', 'bold': None, 'italic': None, 'font_size': 18.0}, {'text': ' created for a word processing formatting project. The purpose of this file is to allow the user to apply formatting such as text justification, font size, font color, line spacing, and paragraph alignment.', 'bold': None, 'italic': None, 'font_size': 18.0}]}, {'text': 'Project Description', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Project Description', 'bold': None, 'italic': None, 'font_size': 10.0}]}, {'text': 'The project focuses on understanding basic document formatting techniques commonly used in professional and academic documents. Students are expected to modify the appearance of the text without changing the actual content provided in this document.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'The project focuses on understanding basic document formatting techniques commonly used in professional and academic documents. Students are expected to modify the appearance of the text without changing the actual content provided in this document.', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Key points:', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'Key points:', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To handle abc', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To handle ', 'bold': None, 'italic': None, 'font_size': None}, {'text': 'abc', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To handle research', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To handle research', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'To ease assignment', 'style': 'Normal', 'alignment': <WD_PARAGRAPH_ALIGNMENT.CENTER: 1>, 'runs': [{'text': 'To ease assignment', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'Conclusion', 'style': 'Heading 1', 'alignment': None, 'runs': [{'text': 'Conclusion', 'bold': None, 'italic': None, 'font_size': None}]}, {'text': 'By applying appropriate formatting styles, this document can be transformed into a well-structured and visually appealing report. Proper formatting improves readability and enhances the overall presentation of written content.', 'style': 'Normal', 'alignment': None, 'runs': [{'text': 'By applying appropriate formatting styles, this document can be transformed into a well-structured and visually appealing report. Proper formatting improves readability and enhances the overall presentation of written content.', 'bold': None, 'italic': None, 'font_size': 14.0}]}]\n",
      "document_type='academic' headings={'Introduction': 1, 'Project Description': 1, 'Conclusion': 1} body_font_size=12 heading_font_size=14 alignment='left' issues_found=['Inconsistent font sizes in body text (e.g., 18.0 in some parts, 14.0 in others)', \"List items ('Qwertyujnbvcddx', 'Ertyhhbvcx', 'sdfgbhcd') are not properly formatted as a list\", \"Some headings have inconsistent font sizes (e.g., 'Project Description' has font_size: 10.0)\", \"Centered text under 'Key points:' may not be appropriate for an academic document\", 'No clear line spacing or paragraph indentation specified']\n"
     ]
    }
   ],
   "source": [
    "# result = app.invoke({\n",
    "#     \"doc\": \"Draft.docx\",\n",
    "#     \"structDoc\": [],\n",
    "#     \"formatPlan\": [],\n",
    "#     \"output_doc\": []\n",
    "# })\n",
    "\n",
    "# print(result[\"output_doc\"])\n",
    "# print(result[\"structDoc\"])\n",
    "# print(result[\"formatPlan\"])\n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 53,
   "id": "281bf8f1",
   "metadata": {},
   "outputs": [
    {
     "ename": "ValueError",
     "evalue": "Invalid input type <class 'langchain_core.prompts.prompt.PromptTemplate'>. Must be a PromptValue, str, or list of BaseMessages.",
     "output_type": "error",
     "traceback": [
      "\u001b[31m---------------------------------------------------------------------------\u001b[39m",
      "\u001b[31mValueError\u001b[39m                                Traceback (most recent call last)",
      "\u001b[36mCell\u001b[39m\u001b[36m \u001b[39m\u001b[32mIn[53]\u001b[39m\u001b[32m, line 1\u001b[39m\n\u001b[32m----> \u001b[39m\u001b[32m1\u001b[39m res = \u001b[43mapp\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m(\u001b[49m\u001b[43m{\u001b[49m\n\u001b[32m      2\u001b[39m \u001b[43m    \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mdoc\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m:\u001b[49m\u001b[43m \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mDraft.docx\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m,\u001b[49m\n\u001b[32m      3\u001b[39m \u001b[43m    \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43muser_instruction\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m:\u001b[49m\u001b[43m \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mMake font size 12 and justify text\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\n\u001b[32m      4\u001b[39m \u001b[43m}\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m      6\u001b[39m \u001b[38;5;28mprint\u001b[39m(res[\u001b[33m\"\u001b[39m\u001b[33moutput_doc\u001b[39m\u001b[33m\"\u001b[39m])\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\pregel\\main.py:3026\u001b[39m, in \u001b[36mPregel.invoke\u001b[39m\u001b[34m(self, input, config, context, stream_mode, print_mode, output_keys, interrupt_before, interrupt_after, durability, **kwargs)\u001b[39m\n\u001b[32m   3023\u001b[39m chunks: \u001b[38;5;28mlist\u001b[39m[\u001b[38;5;28mdict\u001b[39m[\u001b[38;5;28mstr\u001b[39m, Any] | Any] = []\n\u001b[32m   3024\u001b[39m interrupts: \u001b[38;5;28mlist\u001b[39m[Interrupt] = []\n\u001b[32m-> \u001b[39m\u001b[32m3026\u001b[39m \u001b[43m\u001b[49m\u001b[38;5;28;43;01mfor\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mchunk\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;129;43;01min\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43mstream\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m   3027\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;28;43minput\u001b[39;49m\u001b[43m,\u001b[49m\n\u001b[32m   3028\u001b[39m \u001b[43m    \u001b[49m\u001b[43mconfig\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3029\u001b[39m \u001b[43m    \u001b[49m\u001b[43mcontext\u001b[49m\u001b[43m=\u001b[49m\u001b[43mcontext\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3030\u001b[39m \u001b[43m    \u001b[49m\u001b[43mstream_mode\u001b[49m\u001b[43m=\u001b[49m\u001b[43m[\u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mupdates\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mvalues\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m]\u001b[49m\n\u001b[32m   3031\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;28;43;01mif\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mstream_mode\u001b[49m\u001b[43m \u001b[49m\u001b[43m==\u001b[49m\u001b[43m \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mvalues\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\n\u001b[32m   3032\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;28;43;01melse\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mstream_mode\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3033\u001b[39m \u001b[43m    \u001b[49m\u001b[43mprint_mode\u001b[49m\u001b[43m=\u001b[49m\u001b[43mprint_mode\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3034\u001b[39m \u001b[43m    \u001b[49m\u001b[43moutput_keys\u001b[49m\u001b[43m=\u001b[49m\u001b[43moutput_keys\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3035\u001b[39m \u001b[43m    \u001b[49m\u001b[43minterrupt_before\u001b[49m\u001b[43m=\u001b[49m\u001b[43minterrupt_before\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3036\u001b[39m \u001b[43m    \u001b[49m\u001b[43minterrupt_after\u001b[49m\u001b[43m=\u001b[49m\u001b[43minterrupt_after\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3037\u001b[39m \u001b[43m    \u001b[49m\u001b[43mdurability\u001b[49m\u001b[43m=\u001b[49m\u001b[43mdurability\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3038\u001b[39m \u001b[43m    \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   3039\u001b[39m \u001b[43m\u001b[49m\u001b[43m)\u001b[49m\u001b[43m:\u001b[49m\n\u001b[32m   3040\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;28;43;01mif\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mstream_mode\u001b[49m\u001b[43m \u001b[49m\u001b[43m==\u001b[49m\u001b[43m \u001b[49m\u001b[33;43m\"\u001b[39;49m\u001b[33;43mvalues\u001b[39;49m\u001b[33;43m\"\u001b[39;49m\u001b[43m:\u001b[49m\n\u001b[32m   3041\u001b[39m \u001b[43m        \u001b[49m\u001b[38;5;28;43;01mif\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[38;5;28;43mlen\u001b[39;49m\u001b[43m(\u001b[49m\u001b[43mchunk\u001b[49m\u001b[43m)\u001b[49m\u001b[43m \u001b[49m\u001b[43m==\u001b[49m\u001b[43m \u001b[49m\u001b[32;43m2\u001b[39;49m\u001b[43m:\u001b[49m\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\pregel\\main.py:2647\u001b[39m, in \u001b[36mPregel.stream\u001b[39m\u001b[34m(self, input, config, context, stream_mode, print_mode, output_keys, interrupt_before, interrupt_after, durability, subgraphs, debug, **kwargs)\u001b[39m\n\u001b[32m   2645\u001b[39m \u001b[38;5;28;01mfor\u001b[39;00m task \u001b[38;5;129;01min\u001b[39;00m loop.match_cached_writes():\n\u001b[32m   2646\u001b[39m     loop.output_writes(task.id, task.writes, cached=\u001b[38;5;28;01mTrue\u001b[39;00m)\n\u001b[32m-> \u001b[39m\u001b[32m2647\u001b[39m \u001b[43m\u001b[49m\u001b[38;5;28;43;01mfor\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43m_\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;129;43;01min\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mrunner\u001b[49m\u001b[43m.\u001b[49m\u001b[43mtick\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m   2648\u001b[39m \u001b[43m    \u001b[49m\u001b[43m[\u001b[49m\u001b[43mt\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;28;43;01mfor\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mt\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;129;43;01min\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mloop\u001b[49m\u001b[43m.\u001b[49m\u001b[43mtasks\u001b[49m\u001b[43m.\u001b[49m\u001b[43mvalues\u001b[49m\u001b[43m(\u001b[49m\u001b[43m)\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;28;43;01mif\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[38;5;129;43;01mnot\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43mt\u001b[49m\u001b[43m.\u001b[49m\u001b[43mwrites\u001b[49m\u001b[43m]\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   2649\u001b[39m \u001b[43m    \u001b[49m\u001b[43mtimeout\u001b[49m\u001b[43m=\u001b[49m\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43mstep_timeout\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   2650\u001b[39m \u001b[43m    \u001b[49m\u001b[43mget_waiter\u001b[49m\u001b[43m=\u001b[49m\u001b[43mget_waiter\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   2651\u001b[39m \u001b[43m    \u001b[49m\u001b[43mschedule_task\u001b[49m\u001b[43m=\u001b[49m\u001b[43mloop\u001b[49m\u001b[43m.\u001b[49m\u001b[43maccept_push\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   2652\u001b[39m \u001b[43m\u001b[49m\u001b[43m)\u001b[49m\u001b[43m:\u001b[49m\n\u001b[32m   2653\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;66;43;03m# emit output\u001b[39;49;00m\n\u001b[32m   2654\u001b[39m \u001b[43m    \u001b[49m\u001b[38;5;28;43;01myield from\u001b[39;49;00m\u001b[43m \u001b[49m\u001b[43m_output\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m   2655\u001b[39m \u001b[43m        \u001b[49m\u001b[43mstream_mode\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mprint_mode\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43msubgraphs\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mstream\u001b[49m\u001b[43m.\u001b[49m\u001b[43mget\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mqueue\u001b[49m\u001b[43m.\u001b[49m\u001b[43mEmpty\u001b[49m\n\u001b[32m   2656\u001b[39m \u001b[43m    \u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m   2657\u001b[39m loop.after_tick()\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\pregel\\_runner.py:162\u001b[39m, in \u001b[36mPregelRunner.tick\u001b[39m\u001b[34m(self, tasks, reraise, timeout, retry_policy, get_waiter, schedule_task)\u001b[39m\n\u001b[32m    160\u001b[39m t = tasks[\u001b[32m0\u001b[39m]\n\u001b[32m    161\u001b[39m \u001b[38;5;28;01mtry\u001b[39;00m:\n\u001b[32m--> \u001b[39m\u001b[32m162\u001b[39m     \u001b[43mrun_with_retry\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m    163\u001b[39m \u001b[43m        \u001b[49m\u001b[43mt\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    164\u001b[39m \u001b[43m        \u001b[49m\u001b[43mretry_policy\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    165\u001b[39m \u001b[43m        \u001b[49m\u001b[43mconfigurable\u001b[49m\u001b[43m=\u001b[49m\u001b[43m{\u001b[49m\n\u001b[32m    166\u001b[39m \u001b[43m            \u001b[49m\u001b[43mCONFIG_KEY_CALL\u001b[49m\u001b[43m:\u001b[49m\u001b[43m \u001b[49m\u001b[43mpartial\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m    167\u001b[39m \u001b[43m                \u001b[49m\u001b[43m_call\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    168\u001b[39m \u001b[43m                \u001b[49m\u001b[43mweakref\u001b[49m\u001b[43m.\u001b[49m\u001b[43mref\u001b[49m\u001b[43m(\u001b[49m\u001b[43mt\u001b[49m\u001b[43m)\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    169\u001b[39m \u001b[43m                \u001b[49m\u001b[43mretry_policy\u001b[49m\u001b[43m=\u001b[49m\u001b[43mretry_policy\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    170\u001b[39m \u001b[43m                \u001b[49m\u001b[43mfutures\u001b[49m\u001b[43m=\u001b[49m\u001b[43mweakref\u001b[49m\u001b[43m.\u001b[49m\u001b[43mref\u001b[49m\u001b[43m(\u001b[49m\u001b[43mfutures\u001b[49m\u001b[43m)\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    171\u001b[39m \u001b[43m                \u001b[49m\u001b[43mschedule_task\u001b[49m\u001b[43m=\u001b[49m\u001b[43mschedule_task\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    172\u001b[39m \u001b[43m                \u001b[49m\u001b[43msubmit\u001b[49m\u001b[43m=\u001b[49m\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43msubmit\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    173\u001b[39m \u001b[43m            \u001b[49m\u001b[43m)\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    174\u001b[39m \u001b[43m        \u001b[49m\u001b[43m}\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m    175\u001b[39m \u001b[43m    \u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m    176\u001b[39m     \u001b[38;5;28mself\u001b[39m.commit(t, \u001b[38;5;28;01mNone\u001b[39;00m)\n\u001b[32m    177\u001b[39m \u001b[38;5;28;01mexcept\u001b[39;00m \u001b[38;5;167;01mException\u001b[39;00m \u001b[38;5;28;01mas\u001b[39;00m exc:\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\pregel\\_retry.py:42\u001b[39m, in \u001b[36mrun_with_retry\u001b[39m\u001b[34m(task, retry_policy, configurable)\u001b[39m\n\u001b[32m     40\u001b[39m     task.writes.clear()\n\u001b[32m     41\u001b[39m     \u001b[38;5;66;03m# run the task\u001b[39;00m\n\u001b[32m---> \u001b[39m\u001b[32m42\u001b[39m     \u001b[38;5;28;01mreturn\u001b[39;00m \u001b[43mtask\u001b[49m\u001b[43m.\u001b[49m\u001b[43mproc\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m(\u001b[49m\u001b[43mtask\u001b[49m\u001b[43m.\u001b[49m\u001b[43minput\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mconfig\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m     43\u001b[39m \u001b[38;5;28;01mexcept\u001b[39;00m ParentCommand \u001b[38;5;28;01mas\u001b[39;00m exc:\n\u001b[32m     44\u001b[39m     ns: \u001b[38;5;28mstr\u001b[39m = config[CONF][CONFIG_KEY_CHECKPOINT_NS]\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\_internal\\_runnable.py:657\u001b[39m, in \u001b[36mRunnableSeq.invoke\u001b[39m\u001b[34m(self, input, config, **kwargs)\u001b[39m\n\u001b[32m    655\u001b[39m     \u001b[38;5;66;03m# run in context\u001b[39;00m\n\u001b[32m    656\u001b[39m     \u001b[38;5;28;01mwith\u001b[39;00m set_config_context(config, run) \u001b[38;5;28;01mas\u001b[39;00m context:\n\u001b[32m--> \u001b[39m\u001b[32m657\u001b[39m         \u001b[38;5;28minput\u001b[39m = \u001b[43mcontext\u001b[49m\u001b[43m.\u001b[49m\u001b[43mrun\u001b[49m\u001b[43m(\u001b[49m\u001b[43mstep\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[38;5;28;43minput\u001b[39;49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mconfig\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m    658\u001b[39m \u001b[38;5;28;01melse\u001b[39;00m:\n\u001b[32m    659\u001b[39m     \u001b[38;5;28minput\u001b[39m = step.invoke(\u001b[38;5;28minput\u001b[39m, config)\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langgraph\\_internal\\_runnable.py:401\u001b[39m, in \u001b[36mRunnableCallable.invoke\u001b[39m\u001b[34m(self, input, config, **kwargs)\u001b[39m\n\u001b[32m    399\u001b[39m         run_manager.on_chain_end(ret)\n\u001b[32m    400\u001b[39m \u001b[38;5;28;01melse\u001b[39;00m:\n\u001b[32m--> \u001b[39m\u001b[32m401\u001b[39m     ret = \u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43mfunc\u001b[49m\u001b[43m(\u001b[49m\u001b[43m*\u001b[49m\u001b[43margs\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m    402\u001b[39m \u001b[38;5;28;01mif\u001b[39;00m \u001b[38;5;28mself\u001b[39m.recurse \u001b[38;5;129;01mand\u001b[39;00m \u001b[38;5;28misinstance\u001b[39m(ret, Runnable):\n\u001b[32m    403\u001b[39m     \u001b[38;5;28;01mreturn\u001b[39;00m ret.invoke(\u001b[38;5;28minput\u001b[39m, config)\n",
      "\u001b[36mCell\u001b[39m\u001b[36m \u001b[39m\u001b[32mIn[47]\u001b[39m\u001b[32m, line 12\u001b[39m, in \u001b[36mchatCommandNode\u001b[39m\u001b[34m(state)\u001b[39m\n\u001b[32m      2\u001b[39m instruction = state[\u001b[33m\"\u001b[39m\u001b[33muser_instruction\u001b[39m\u001b[33m\"\u001b[39m]\n\u001b[32m      4\u001b[39m prompt = PromptTemplate(\n\u001b[32m      5\u001b[39m     template=\u001b[33m\"\"\"\u001b[39m\n\u001b[32m      6\u001b[39m \u001b[33m    Convert the user instruction into a formatting command.\u001b[39m\n\u001b[32m   (...)\u001b[39m\u001b[32m      9\u001b[39m     input_variables=[\u001b[33m\"\u001b[39m\u001b[33minstruction\u001b[39m\u001b[33m\"\u001b[39m]\n\u001b[32m     10\u001b[39m )\n\u001b[32m---> \u001b[39m\u001b[32m12\u001b[39m command = \u001b[43mmodel\u001b[49m\u001b[43m.\u001b[49m\u001b[43mwith_structured_output\u001b[49m\u001b[43m(\u001b[49m\u001b[43mFormatCommand\u001b[49m\u001b[43m)\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m(\u001b[49m\u001b[43mprompt\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m     13\u001b[39m state[\u001b[33m\"\u001b[39m\u001b[33mformatCmd\u001b[39m\u001b[33m\"\u001b[39m] = command\n\u001b[32m     14\u001b[39m \u001b[38;5;28;01mreturn\u001b[39;00m state\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langchain_core\\runnables\\base.py:3047\u001b[39m, in \u001b[36mRunnableSequence.invoke\u001b[39m\u001b[34m(self, input, config, **kwargs)\u001b[39m\n\u001b[32m   3045\u001b[39m \u001b[38;5;28;01mwith\u001b[39;00m set_config_context(config) \u001b[38;5;28;01mas\u001b[39;00m context:\n\u001b[32m   3046\u001b[39m     \u001b[38;5;28;01mif\u001b[39;00m i == \u001b[32m0\u001b[39m:\n\u001b[32m-> \u001b[39m\u001b[32m3047\u001b[39m         input_ = \u001b[43mcontext\u001b[49m\u001b[43m.\u001b[49m\u001b[43mrun\u001b[49m\u001b[43m(\u001b[49m\u001b[43mstep\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43minput_\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43mconfig\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m)\u001b[49m\n\u001b[32m   3048\u001b[39m     \u001b[38;5;28;01melse\u001b[39;00m:\n\u001b[32m   3049\u001b[39m         input_ = context.run(step.invoke, input_, config)\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langchain_core\\runnables\\base.py:5441\u001b[39m, in \u001b[36mRunnableBindingBase.invoke\u001b[39m\u001b[34m(self, input, config, **kwargs)\u001b[39m\n\u001b[32m   5434\u001b[39m \u001b[38;5;129m@override\u001b[39m\n\u001b[32m   5435\u001b[39m \u001b[38;5;28;01mdef\u001b[39;00m\u001b[38;5;250m \u001b[39m\u001b[34minvoke\u001b[39m(\n\u001b[32m   5436\u001b[39m     \u001b[38;5;28mself\u001b[39m,\n\u001b[32m   (...)\u001b[39m\u001b[32m   5439\u001b[39m     **kwargs: Optional[Any],\n\u001b[32m   5440\u001b[39m ) -> Output:\n\u001b[32m-> \u001b[39m\u001b[32m5441\u001b[39m     \u001b[38;5;28;01mreturn\u001b[39;00m \u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43mbound\u001b[49m\u001b[43m.\u001b[49m\u001b[43minvoke\u001b[49m\u001b[43m(\u001b[49m\n\u001b[32m   5442\u001b[39m \u001b[43m        \u001b[49m\u001b[38;5;28;43minput\u001b[39;49m\u001b[43m,\u001b[49m\n\u001b[32m   5443\u001b[39m \u001b[43m        \u001b[49m\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43m_merge_configs\u001b[49m\u001b[43m(\u001b[49m\u001b[43mconfig\u001b[49m\u001b[43m)\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   5444\u001b[39m \u001b[43m        \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43m{\u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m,\u001b[49m\u001b[43m \u001b[49m\u001b[43m*\u001b[49m\u001b[43m*\u001b[49m\u001b[43mkwargs\u001b[49m\u001b[43m}\u001b[49m\u001b[43m,\u001b[49m\n\u001b[32m   5445\u001b[39m \u001b[43m    \u001b[49m\u001b[43m)\u001b[49m\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langchain_core\\language_models\\chat_models.py:384\u001b[39m, in \u001b[36mBaseChatModel.invoke\u001b[39m\u001b[34m(self, input, config, stop, **kwargs)\u001b[39m\n\u001b[32m    371\u001b[39m \u001b[38;5;129m@override\u001b[39m\n\u001b[32m    372\u001b[39m \u001b[38;5;28;01mdef\u001b[39;00m\u001b[38;5;250m \u001b[39m\u001b[34minvoke\u001b[39m(\n\u001b[32m    373\u001b[39m     \u001b[38;5;28mself\u001b[39m,\n\u001b[32m   (...)\u001b[39m\u001b[32m    378\u001b[39m     **kwargs: Any,\n\u001b[32m    379\u001b[39m ) -> BaseMessage:\n\u001b[32m    380\u001b[39m     config = ensure_config(config)\n\u001b[32m    381\u001b[39m     \u001b[38;5;28;01mreturn\u001b[39;00m cast(\n\u001b[32m    382\u001b[39m         \u001b[33m\"\u001b[39m\u001b[33mChatGeneration\u001b[39m\u001b[33m\"\u001b[39m,\n\u001b[32m    383\u001b[39m         \u001b[38;5;28mself\u001b[39m.generate_prompt(\n\u001b[32m--> \u001b[39m\u001b[32m384\u001b[39m             [\u001b[38;5;28;43mself\u001b[39;49m\u001b[43m.\u001b[49m\u001b[43m_convert_input\u001b[49m\u001b[43m(\u001b[49m\u001b[38;5;28;43minput\u001b[39;49m\u001b[43m)\u001b[49m],\n\u001b[32m    385\u001b[39m             stop=stop,\n\u001b[32m    386\u001b[39m             callbacks=config.get(\u001b[33m\"\u001b[39m\u001b[33mcallbacks\u001b[39m\u001b[33m\"\u001b[39m),\n\u001b[32m    387\u001b[39m             tags=config.get(\u001b[33m\"\u001b[39m\u001b[33mtags\u001b[39m\u001b[33m\"\u001b[39m),\n\u001b[32m    388\u001b[39m             metadata=config.get(\u001b[33m\"\u001b[39m\u001b[33mmetadata\u001b[39m\u001b[33m\"\u001b[39m),\n\u001b[32m    389\u001b[39m             run_name=config.get(\u001b[33m\"\u001b[39m\u001b[33mrun_name\u001b[39m\u001b[33m\"\u001b[39m),\n\u001b[32m    390\u001b[39m             run_id=config.pop(\u001b[33m\"\u001b[39m\u001b[33mrun_id\u001b[39m\u001b[33m\"\u001b[39m, \u001b[38;5;28;01mNone\u001b[39;00m),\n\u001b[32m    391\u001b[39m             **kwargs,\n\u001b[32m    392\u001b[39m         ).generations[\u001b[32m0\u001b[39m][\u001b[32m0\u001b[39m],\n\u001b[32m    393\u001b[39m     ).message\n",
      "\u001b[36mFile \u001b[39m\u001b[32mc:\\Users\\oshik\\.vscode\\ProgrammingFolder\\Ai\\langgraph\\env\\Lib\\site-packages\\langchain_core\\language_models\\chat_models.py:369\u001b[39m, in \u001b[36mBaseChatModel._convert_input\u001b[39m\u001b[34m(self, model_input)\u001b[39m\n\u001b[32m    364\u001b[39m     \u001b[38;5;28;01mreturn\u001b[39;00m ChatPromptValue(messages=convert_to_messages(model_input))\n\u001b[32m    365\u001b[39m msg = (\n\u001b[32m    366\u001b[39m     \u001b[33mf\u001b[39m\u001b[33m\"\u001b[39m\u001b[33mInvalid input type \u001b[39m\u001b[38;5;132;01m{\u001b[39;00m\u001b[38;5;28mtype\u001b[39m(model_input)\u001b[38;5;132;01m}\u001b[39;00m\u001b[33m. \u001b[39m\u001b[33m\"\u001b[39m\n\u001b[32m    367\u001b[39m     \u001b[33m\"\u001b[39m\u001b[33mMust be a PromptValue, str, or list of BaseMessages.\u001b[39m\u001b[33m\"\u001b[39m\n\u001b[32m    368\u001b[39m )\n\u001b[32m--> \u001b[39m\u001b[32m369\u001b[39m \u001b[38;5;28;01mraise\u001b[39;00m \u001b[38;5;167;01mValueError\u001b[39;00m(msg)\n",
      "\u001b[31mValueError\u001b[39m: Invalid input type <class 'langchain_core.prompts.prompt.PromptTemplate'>. Must be a PromptValue, str, or list of BaseMessages.",
      "During task with name 'chat_command' and id '027c69e5-2bbf-01c5-e168-6eb8c4059c5d'"
     ]
    }
   ],
   "source": [
    "res = app.invoke({\n",
    "    \"doc\": \"Draft.docx\",\n",
    "    \"user_instruction\": \"Make font size 12 and justify text\"\n",
    "})\n",
    "\n",
    "print(res[\"output_doc\"])"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "id": "ee5c2187",
   "metadata": {},
   "outputs": [],
   "source": []
  }
 ],
 "metadata": {
  "kernelspec": {
   "display_name": "env",
   "language": "python",
   "name": "python3"
  },
  "language_info": {
   "codemirror_mode": {
    "name": "ipython",
    "version": 3
   },
   "file_extension": ".py",
   "mimetype": "text/x-python",
   "name": "python",
   "nbconvert_exporter": "python",
   "pygments_lexer": "ipython3",
   "version": "3.12.5"
  }
 },
 "nbformat": 4,
 "nbformat_minor": 5
}
