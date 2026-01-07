"""
Word Document Formatter - Flask Backend
Converts Word documents with AI-powered formatting analysis and commands.
"""

import os
import json
import mammoth
from typing import TypedDict, Annotated, List, Dict, Optional, Literal, Union
from enum import Enum

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import tempfile
import uuid
from langgraph.graph import add_messages
from langgraph.graph import StateGraph
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field



from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_utils import docx_to_html



# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Initialize LLM model
model = ChatOpenAI(
    model="mistralai/devstral-2512:free",
    openai_api_base="https://openrouter.ai/api/v1",
    openai_api_key= os.getenv("OPENAI_API_KEY", ""),
)

# Setting Directory
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================================
# Pydantic Models and Enums
# ============================================================================

class DocStructure(BaseModel):
    """Document structure analysis result."""
    document_type: str = Field(
        description="Type of document: academic, research, business, general"
    )
    headings: Dict[str, int] = Field(
        description="Detected headings mapped to heading level (e.g., {'Introduction': 1})"
    )
    body_font_size: int = Field(
        description="Target font size for body text"
    )
    heading_font_size: int = Field(
        description="Target font size for headings"
    )
    alignment: str = Field(
        description="Text alignment: left, right, center, justify"
    )
    issues_found: List[str] = Field(
        description="Formatting or structural issues detected"
    )


class ActionEnum(str, Enum):
    """Available formatting actions."""
    SET_FONT_SIZE = "set_font_size"
    SET_ALIGNMENT = "set_alignment"
    ALIGN_TEXT = "align_text"
    FORMAT_AS_LIST = "format_as_list"
    REMOVE_EMPTY_PARAGRAPHS = "remove_empty_paragraphs"


class FormatCommand(BaseModel):
    """Single formatting command."""
    action: ActionEnum
    target: str
    value: Union[int, str]
    source: Literal["auto", "user"]


class FormatCommandList(BaseModel):
    """List of formatting commands."""
    commands: List[FormatCommand]


# ============================================================================
# State Management
# ============================================================================

def last_write(_, new):
    """Reducer function for state updates."""
    return new


class WordFormatter(TypedDict, total=False):
    """State graph type definition."""
    doc: str
    document: Document
    structDoc: Annotated[list, last_write]
    AutoDetect: Annotated[DocStructure, last_write]
    AutoDetectCmd: Annotated[list, last_write]
    userCmd: Annotated[list, last_write]
    formatCmd: Annotated[list, last_write]
    output_doc: Optional[str]
    output_pdf: Optional[str]
    user_instruction: str


# ============================================================================
# Document Parsing Functions
# ============================================================================

def load_docx(path: str) -> List[Dict]:
    """
    Parse a Word document into structured data.
    
    Args:
        path: Path to the .docx file
        
    Returns:
        List of paragraph dictionaries with text, style, alignment, and run info
    """
    doc = Document(path)
    paragraphs = []

    for p in doc.paragraphs:
        paragraphs.append({
            "text": p.text.strip(),
            "style": p.style.name if p.style else None,
            "alignment": p.alignment,
            "runs": [
                {
                    "text": r.text,
                    "bold": r.bold,
                    "italic": r.italic,
                    "font_size": r.font.size.pt if r.font.size else None
                }
                for r in p.runs
            ]
        })
    return paragraphs


# ============================================================================
# Constants
# ============================================================================

ACTION_DOC = """
Allowed actions (field `action`):
- "set_font_size": Change font size for the specified target.
- "set_alignment": Change paragraph alignment.
- "align_text": Align only paragraphs containing specific text.
- "format_as_list": Convert paragraphs into a bulleted list.
- "remove_empty_paragraphs": Remove trailing empty paragraphs.


Allowed target (field `target`):
- anything present in the document can also be the target
- "heading": All heading-style paragraphs (Heading 1–3).
- "subheading": Subsections under headings.
- "body_text": Normal body paragraphs.
- "key_points": Paragraphs marked or inferred as key points.
- "list_items": Existing bullet or numbered list items.
- "document": Entire document.
- "end_of_document": Only trailing paragraphs at the end.


Return float values as integer not as string
Do NOT invent new actions like "justify", "bolden", etc.
"""


# ============================================================================
# LangGraph Node Functions
# ============================================================================

def load_doc_node(state: WordFormatter) -> dict:
    """Node: Load document and parse structure."""
    doc = Document(state["doc"])
    return {
        "document": doc,
        "structDoc": load_docx(state["doc"])
    }


def doc_structure_node(state: WordFormatter) -> dict:
    """Node: Analyze document structure using LLM."""
    struct_doc = state["structDoc"]
    parser = PydanticOutputParser(pydantic_object=DocStructure)

    prompt = PromptTemplate(
        template="""
You are a document structure analysis agent.

Analyze the following document structure and produce a formatting plan.

Document structure:
{struct_doc}

{format_instructions}
""",
        input_variables=["struct_doc"],
        partial_variables={
            "format_instructions": parser.get_format_instructions()
        }
    )

    response = model.invoke(
        prompt.format(struct_doc=struct_doc)
    )

    return {
        "AutoDetect": parser.parse(response.content)
    }


def auto_detect_command_node(state: WordFormatter) -> dict:
    """Node: Convert detected issues into formatting commands."""
    instruction = state["AutoDetect"]
    cmd_parser = PydanticOutputParser(pydantic_object=FormatCommandList)

    prompt = PromptTemplate(
        template="""
You are a document formatting planner.

Given these issues:
{issues}
{action_doc}

Generate formatting commands.
Return them strictly in this format:
{format_instructions}
""",
        input_variables=["issues"],
        partial_variables={
            "format_instructions": cmd_parser.get_format_instructions(),
            "action_doc": ACTION_DOC,
        }
    )

    response = model.invoke(prompt.format(issues=instruction.issues_found))
    parsed = cmd_parser.parse(response.content)
    print(parsed.commands)
    return {"AutoDetectCmd": parsed.commands}


def chat_command_node(state: WordFormatter) -> dict:
    """Node: Convert user instructions into formatting commands."""
    instruction = state["user_instruction"]
    cmd_parser = PydanticOutputParser(pydantic_object=FormatCommandList)

    prompt = PromptTemplate(
        template="""
        You are a document formatting planner
        Given the user instruction
        Instruction: {instruction}
        {action_doc}
        Generate formatting commands.
        Return them strictly in this format:
        {format_instructions}
        """,
        input_variables=["instruction"],
        partial_variables={
            "format_instructions": cmd_parser.get_format_instructions(),
            "action_doc": ACTION_DOC,
        }
    )
    
    response = model.invoke(prompt.format(instruction=instruction))
    parsed = cmd_parser.parse(response.content)
    print("chat,", parsed.commands)
    return {
        "userCmd": parsed.commands
    }


def merge_command_node(state: WordFormatter) -> dict:
    """Node: Merge auto-detected and user commands with conflict resolution."""
    TARGET_SCOPE = {
        "document": 0,      # global
        "text": 0,
        "body_text": 1,
        "heading": 2,
        "list_items": 2,
    }
    
    auto_cmd = state.get("AutoDetectCmd") or []
    user_cmd = state.get("userCmd") or []

    if not isinstance(auto_cmd, list):
        raise TypeError(f"AutoDetectCmd must be list, got {type(auto_cmd)}")

    if not isinstance(user_cmd, list):
        raise TypeError(f"userCmd must be list, got {type(user_cmd)}")

    resolved = {}

    def scope(t):
        return TARGET_SCOPE.get(t, 0)

    def should_override(existing, incoming):
        if incoming.source == "user" and existing.source == "auto":
            return True
        return scope(incoming.target) >= scope(existing.target)

    for cmd in auto_cmd + user_cmd:
        key = cmd.action
        if key not in resolved or should_override(resolved[key], cmd):
            resolved[key] = cmd

    print("fm", list(resolved.values()))
    return {"formatCmd": list(resolved.values())}


# ============================================================================
# Formatting Tools
# ============================================================================
def set_font_size_tool(doc: Document, target: str, value: int):
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""

        is_heading = style.startswith("Heading") or style == "Title"
        is_body = style in ("Normal", "Body Text", "Default Paragraph Font")

        if target == "heading" and is_heading:
            for r in p.runs:
                r.font.size = Pt(int(value))

        elif target == "body_text" and is_body:
            for r in p.runs:
                r.font.size = Pt(int(value))

        elif target == "document":
            for r in p.runs:
                r.font.size = Pt(int(value))



def set_alignment_tool(doc: Document, target: str, value: str):
    """Apply alignment formatting to document."""
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    for p in doc.paragraphs:
        if target == "document" or target.lower() in p.text.lower():
            p.paragraph_format.alignment = align_map.get(
                value.lower(), 
                p.paragraph_format.alignment
            )
    return f"Set {target} alignment to {value}"


TOOL_REGISTRY = {
    "set_font_size": set_font_size_tool,
    "set_alignment": set_alignment_tool,
    "align_text": set_alignment_tool  # Alias
}


def tool_cmd_node(state: WordFormatter) -> dict:
    """Node: Apply formatting commands to document."""
    doc = state["document"]
    commands = state.get("formatCmd", [])
    
    applied_actions = []
    skipped_actions = []
    
    for cmd in commands:
        func = TOOL_REGISTRY.get(cmd.action.value)
        if func:
            func(doc, cmd.target, cmd.value)
            applied_actions.append({
                "action": cmd.action.value,
                "target": cmd.target,
                "value": str(cmd.value)
            })
        else:
            skipped_actions.append({
                "action": cmd.action.value,
                "target": cmd.target,
                "reason": "No tool implementation found"
            })



    # Generate output filename in current directory
    output_filename = f"Final_Formatted_{uuid.uuid4().hex[:8]}.docx"
    # output_path = os.path.join(os.getcwd(), output_filename)
    # doc.save(output_path)
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    
    return {
        "output_doc": output_filename,  # Return just filename for download endpoint
        "applied_actions": applied_actions,
        "skipped_actions": skipped_actions
    }


# ============================================================================
# LangGraph Setup
# ============================================================================

def create_graph(checkpointer=None):
    """
    Create and compile the LangGraph workflow.
    
    Args:
        checkpointer: Optional checkpoint saver for state persistence.
                     If None, graph runs without checkpointing.
    """
    graph = StateGraph(WordFormatter)

    # Add nodes
    graph.add_node("load_docx", load_doc_node)
    graph.add_node("doc_structure", doc_structure_node)
    graph.add_node("chat_command", chat_command_node)
    graph.add_node("AutoDetect_command", auto_detect_command_node)
    graph.add_node("merge_command", merge_command_node)
    graph.add_node("Toolcmd", tool_cmd_node)

    # Set entry point
    graph.set_entry_point("load_docx")

    # Add edges
    graph.add_edge("load_docx", "doc_structure")
    graph.add_edge("load_docx", "chat_command")
    graph.add_edge("doc_structure", "AutoDetect_command")

    graph.add_edge("chat_command", "merge_command")
    graph.add_edge("AutoDetect_command", "merge_command")

    graph.add_edge("merge_command", "Toolcmd")
    graph.set_finish_point("Toolcmd")

    # Compile with optional checkpointer
    return graph.compile(checkpointer=checkpointer)


# Initialize graph
graph_app = create_graph()






# ============================================================================
# Flask API Endpoints
# ============================================================================

@app.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint."""
    return jsonify({"status": "healthy", "service": "word-formatter"})


@app.route("/format-document", methods=["POST"])
def format_document():
    """
    Main endpoint for document formatting.
    
    Accepts either:
    1. File upload (multipart/form-data) with 'file' field
    2. JSON with doc_path field
    
    JSON payload (alternative):
    {
        "doc_path": "path/to/document.docx",
        "user_instruction": "Justify the document"  # Optional
    }
    
    Returns:
    {
        "success": true,
        "output_doc": "path/to/output.docx",
        "format_cmd": [...],
        "auto_detect": {...},
        "auto_detect_cmd": [...],
        "user_cmd": [...],
        "applied_actions": [...],
        "skipped_actions": [...]
    }
    """
    try:
        doc_path = None
        user_instruction = ""
        temp_file = None
        
        # Check if file was uploaded
        if 'file' in request.files:
            uploaded_file = request.files['file']
            if uploaded_file.filename == '':
                return jsonify({
                    "success": False,
                    "error": "No file selected"
                }), 400
            
            # Save uploaded file to temporary location
            file_ext = os.path.splitext(uploaded_file.filename)[1]
            if file_ext.lower() != '.docx':
                return jsonify({
                    "success": False,
                    "error": "Only .docx files are supported"
                }), 400
            
            # Create temp file
            temp_dir = tempfile.gettempdir()
            temp_filename = f"{uuid.uuid4()}{file_ext}"
            temp_file = os.path.join(temp_dir, temp_filename)
            uploaded_file.save(temp_file)
            doc_path = temp_file
            
            # Get user instruction from form data
            user_instruction = request.form.get("user_instruction", "")
        
        # Otherwise, try JSON payload
        elif request.is_json:
            data = request.get_json()
            doc_path = data.get("doc_path")
            user_instruction = data.get("user_instruction", "")
        
        else:
            return jsonify({
                "success": False,
                "error": "Either provide a file upload or JSON with doc_path"
            }), 400
        
        if not doc_path:
            return jsonify({
                "success": False,
                "error": "doc_path or file upload is required"
            }), 400
        
        # Check if file exists
        if not os.path.exists(doc_path):
            return jsonify({
                "success": False,
                "error": f"Document not found: {doc_path}"
            }), 404
        
        # Prepare initial state
        initial_state = {
            "doc": doc_path,
            "user_instruction": user_instruction
        }
        
        # Execute graph
        result = graph_app.invoke(initial_state)
        
        # Serialize result for JSON response
        response_data = {
            "success": True,
            "output_doc": result.get("output_doc"),
            "applied_actions": result.get("applied_actions", []),
            "skipped_actions": result.get("skipped_actions", []),
        }
        
        # Serialize format commands
        if "formatCmd" in result:
            response_data["format_cmd"] = [
                {
                    "action": cmd.action.value,
                    "target": cmd.target,
                    "value": str(cmd.value),
                    "source": cmd.source
                }
                for cmd in result["formatCmd"]
            ]
        
        # Serialize auto detect
        if "AutoDetect" in result:
            auto_detect = result["AutoDetect"]
            response_data["auto_detect"] = {
                "document_type": auto_detect.document_type,
                "headings": auto_detect.headings,
                "body_font_size": auto_detect.body_font_size,
                "heading_font_size": auto_detect.heading_font_size,
                "alignment": auto_detect.alignment,
                "issues_found": auto_detect.issues_found
            }
        
        # Serialize auto detect commands
        if "AutoDetectCmd" in result:
            response_data["auto_detect_cmd"] = [
                {
                    "action": cmd.action.value,
                    "target": cmd.target,
                    "value": str(cmd.value),
                    "source": cmd.source
                }
                for cmd in result["AutoDetectCmd"]
            ]
        
        # Serialize user commands
        if "userCmd" in result:
            response_data["user_cmd"] = [
                {
                    "action": cmd.action.value,
                    "target": cmd.target,
                    "value": str(cmd.value),
                    "source": cmd.source
                }
                for cmd in result["userCmd"]
            ]
        
        # Clean up temp file if it was uploaded
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass
        
        return jsonify(response_data), 200
        
    except Exception as e:
        # Clean up temp file on error
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass
        
        return jsonify({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }), 500
    
@app.route("/preview/<filename>")
def preview_doc(filename):
    path = os.path.join(OUTPUT_DIR, filename)
    return jsonify({"html": docx_to_html(path)})


@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    """Download formatted document."""
    try:
        # Security: only allow downloading from temp directory or current directory
        if not filename.endswith('.docx'):
            return jsonify({"error": "Invalid file type"}), 400
        
        file_path = os.path.join(OUTPUT_DIR, filename)

        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
        
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({
            "error": str(e),
            "error_type": type(e).__name__
        }), 500


# ============================================================================
# Main Entry Point
# ============================================================================

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)

